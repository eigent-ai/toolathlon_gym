"""Evaluation for sf-hr-attrition-word-email."""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_word_doc(agent_workspace, groundtruth_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Attrition_Risk.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "attrition" in p.text.lower() and "risk" in p.text.lower():
            has_heading = True
            break
    check("Document has attrition risk heading", has_heading)

    # Check tables
    check("Document has at least 2 tables", len(doc.tables) >= 2,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 2:
        return False

    # Load groundtruth
    gt_file = os.path.join(groundtruth_workspace, "Attrition_Data.xlsx")
    if not os.path.isfile(gt_file):
        check("Groundtruth file exists", False)
        return False

    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check department table
    gt_dept = list(gt_wb["By Department"].iter_rows(min_row=2, values_only=True))
    table1 = doc.tables[0]
    dept_rows = []
    for row in table1.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        dept_rows.append(cells)

    check("Department table has 7 rows", len(dept_rows) == 7, f"Got {len(dept_rows)} rows")

    for gt_row in gt_dept:
        dept, count, avg_sal, avg_exp = gt_row
        matched = None
        for r in dept_rows:
            if r and dept.lower() in r[0].lower():
                matched = r
                break
        if matched:
            # Check flight risk count
            found_count = False
            for cell in matched[1:]:
                try:
                    val = int(cell.replace(",", ""))
                    if abs(val - count) <= 2:
                        found_count = True
                        break
                except (ValueError, AttributeError):
                    continue
            check(f"Dept {dept} flight risk count", found_count,
                  f"Expected ~{count}")
        else:
            check(f"Dept {dept} found in table", False)

    # Check top risk employees table
    gt_risk = list(gt_wb["Top Risk Employees"].iter_rows(min_row=2, values_only=True))
    table2 = doc.tables[1]
    risk_rows = []
    for row in table2.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        risk_rows.append(cells)

    check("Top risk table has 10 rows", len(risk_rows) == 10, f"Got {len(risk_rows)} rows")

    # Verify at least first employee appears
    if gt_risk and risk_rows:
        first_name = gt_risk[0][1]
        found_first = any(first_name.lower() in " ".join(r).lower() for r in risk_rows)
        check(f"First risk employee '{first_name}' in table", found_first)

    # Check summary
    gt_summary = list(gt_wb["Summary"].iter_rows(min_row=2, values_only=True))
    summary_dict = {r[0]: r[1] for r in gt_summary}
    full_text = " ".join(p.text for p in doc.paragraphs).lower()

    has_total = str(summary_dict.get("Total Employees", "")) in full_text or "50000" in full_text or "50,000" in full_text
    check("Summary mentions total employees", has_total)

    flight_count = str(summary_dict.get("Flight Risk Count", ""))
    has_flight = flight_count in full_text or str(int(flight_count)) in full_text.replace(",", "") if flight_count else False
    check("Summary mentions flight risk count", has_flight)

    return True


def check_email():
    """Check email sent ."""
    print("\n=== Checking Email  ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
        WHERE folder_id != 0
        ORDER BY date DESC
    """)
    emails = cur.fetchall()

    cur.execute("SELECT * FROM email.sent_log ORDER BY id DESC LIMIT 10")
    sent = cur.fetchall()
    cur.close()
    conn.close()

    all_msgs = emails + [(s[2] if len(s) > 2 else None, s[3] if len(s) > 3 else None,
                          s[4] if len(s) > 4 else None, s[5] if len(s) > 5 else None)
                         for s in sent]

    check("At least 1 email sent", len(all_msgs) >= 1,
          f"Found {len(emails)} messages, {len(sent)} sent_log")

    if all_msgs:
        found_attrition = any("attrition" in str(m[0] or "").lower() or "risk" in str(m[0] or "").lower()
                             for m in all_msgs)
        check("Email subject mentions attrition or risk", found_attrition,
              f"Subjects: {[m[0] for m in all_msgs[:5]]}")

        found_recipient = any("hr-director" in str(m[2] or "").lower()
                             for m in all_msgs)
        check("Email sent to hr-director", found_recipient)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("SF HR ATTRITION WORD EMAIL - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace, gt_dir)
    check_email()

    total_pass = PASS_COUNT
    total_fail = FAIL_COUNT
    all_ok = FAIL_COUNT == 0

    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": total_pass, "failed": total_fail, "success": all_ok}, f, indent=2)

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
