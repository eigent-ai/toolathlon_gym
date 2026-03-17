"""Evaluation for sf-support-resolution-word-gcal."""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_word_doc(agent_workspace, groundtruth_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Resolution_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "resolution" in p.text.lower() and "analysis" in p.text.lower():
            has_heading = True
            break
    check("Document has resolution analysis heading", has_heading)

    # Check tables exist
    check("Document has at least 2 tables", len(doc.tables) >= 2,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 2:
        return False

    # Get expected data from DB
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT "PRIORITY", COUNT(*),
               ROUND(AVG(EXTRACT(EPOCH FROM ("RESOLVED_AT" - "CREATED_AT"))/3600)::numeric, 2)
        FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"
        WHERE "STATUS" = 'Resolved' AND "RESOLVED_AT" IS NOT NULL
        GROUP BY "PRIORITY"
        ORDER BY "PRIORITY"
    """)
    priority_data = cur.fetchall()
    conn.close()

    # Check priority table content
    table1 = doc.tables[0]
    rows = []
    for row in table1.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    check("Priority table has 3 rows", len(rows) == 3, f"Got {len(rows)} rows")

    for priority, count, avg_hours in priority_data:
        matched = None
        for r in rows:
            if r and r[0].lower() == priority.lower():
                matched = r
                break
        if matched:
            # Check ticket count
            found_count = False
            for cell in matched[1:]:
                try:
                    val = int(cell.replace(",", ""))
                    if abs(val - count) <= 5:
                        found_count = True
                        break
                except (ValueError, AttributeError):
                    continue
            check(f"Priority {priority} ticket count", found_count,
                  f"Expected ~{count}")

            # Check avg hours
            found_hours = False
            for cell in matched[1:]:
                try:
                    val = float(cell.replace(",", ""))
                    if num_close(val, float(avg_hours), 1.0):
                        found_hours = True
                        break
                except (ValueError, AttributeError):
                    continue
            check(f"Priority {priority} avg hours", found_hours,
                  f"Expected ~{float(avg_hours)}")
        else:
            check(f"Priority {priority} found in table", False)

    # Cross-check with groundtruth
    gt_file = os.path.join(groundtruth_workspace, "Resolution_Data.xlsx")
    if os.path.isfile(gt_file):
        gt_wb = openpyxl.load_workbook(gt_file, data_only=True)
        gt_rows = list(gt_wb["By Priority"].iter_rows(min_row=2, values_only=True))
        for row in gt_rows:
            prio, total, avg_h, min_h, max_h = row
            matched = None
            for r in rows:
                if r and r[0].lower() == prio.lower():
                    matched = r
                    break
            if matched:
                found_avg = False
                for cell in matched[1:]:
                    try:
                        val = float(cell.replace(",", ""))
                        if num_close(val, avg_h, 1.0):
                            found_avg = True
                            break
                    except (ValueError, AttributeError):
                        continue
                check(f"GT cross-check {prio} avg hours", found_avg,
                      f"Expected {avg_h}")

    # Check summary paragraph
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    has_summary = "overall" in full_text or "total" in full_text or "resolved" in full_text
    check("Document has summary text", has_summary)

    return True


def check_gcal():
    """Check Google Calendar event ."""
    print("\n=== Checking Google Calendar  ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, description, start_datetime, end_datetime
        FROM gcal.events
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    check("At least 1 calendar event created", len(events) >= 1,
          f"Found {len(events)}")

    if events:
        found_review = any("resolution" in (e[0] or "").lower() or "review" in (e[0] or "").lower()
                          for e in events)
        check("Event title mentions resolution or review", found_review,
              f"Events: {[e[0] for e in events]}")

        found_date = any("2026-03-12" in str(e[2] or "") for e in events)
        check("Event scheduled for 2026-03-12", found_date,
              f"Dates: {[str(e[2]) for e in events]}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("SF SUPPORT RESOLUTION WORD GCAL - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace, gt_dir)
    check_gcal()

    total_pass = PASS_COUNT
    total_fail = FAIL_COUNT
    all_ok = FAIL_COUNT == 0

    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": total_pass, "failed": total_fail, "success": all_ok}, f, indent=2)

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
