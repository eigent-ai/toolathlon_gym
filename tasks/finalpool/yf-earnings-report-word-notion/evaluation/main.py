"""Evaluation for yf-earnings-report-word-notion."""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_word_doc(agent_workspace, groundtruth_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Earnings_Report.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "earnings" in p.text.lower() and ("analysis" in p.text.lower() or "report" in p.text.lower()):
            has_heading = True
            break
    check("Document has earnings analysis heading", has_heading)

    # Check tables - should have 3 (one per company)
    check("Document has at least 3 tables", len(doc.tables) >= 3,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 3:
        return False

    # Load groundtruth
    gt_file = os.path.join(groundtruth_workspace, "Earnings_Data.xlsx")
    if not os.path.isfile(gt_file):
        check("Groundtruth file exists", False)
        return False

    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    symbols = ["GOOGL", "AMZN", "JNJ"]

    # Check each company section exists
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    for sym in symbols:
        has_section = sym.lower() in full_text
        check(f"Document mentions {sym}", has_section)

    # For each company, check the latest year revenue in some table
    for idx, sym in enumerate(symbols):
        if sym not in gt_wb.sheetnames:
            continue
        gt_rows = list(gt_wb[sym].iter_rows(min_row=2, values_only=True))
        if not gt_rows:
            continue

        latest = gt_rows[-1]
        latest_rev = latest[1]  # Revenue_B

        # Search across all tables for this revenue value
        found_rev = False
        for table in doc.tables:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                for cell in cells:
                    try:
                        val = float(cell.replace(",", "").replace("$", ""))
                        if num_close(val, latest_rev, 2.0):
                            found_rev = True
                            break
                    except (ValueError, AttributeError):
                        continue
                if found_rev:
                    break
            if found_rev:
                break
        check(f"{sym} latest revenue ({latest_rev}B) in document", found_rev)

        latest_ni = latest[2]  # Net_Income_B
        if latest_ni is not None:
            found_ni = False
            for table in doc.tables:
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    for cell in cells:
                        try:
                            val = float(cell.replace(",", "").replace("$", ""))
                            if num_close(val, latest_ni, 2.0):
                                found_ni = True
                                break
                        except (ValueError, AttributeError):
                            continue
                    if found_ni:
                        break
                if found_ni:
                    break
            check(f"{sym} latest net income ({latest_ni}B) in document", found_ni)

    # Check summary
    has_summary = any("summary" in p.text.lower() or "revenue" in p.text.lower()
                      for p in doc.paragraphs if len(p.text) > 50)
    check("Document has summary paragraph", has_summary)

    return True


def check_notion():
    """Check Notion page ."""
    print("\n=== Checking Notion  ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("SELECT id, properties FROM notion.pages")
    pages = cur.fetchall()
    cur.execute("SELECT parent_id, type, block_data FROM notion.blocks")
    blocks = cur.fetchall()
    cur.close()
    conn.close()

    check("At least 1 Notion page created", len(pages) >= 1,
          f"Found {len(pages)}")

    if pages:
        # Check page title
        found_earnings = False
        for page_id, props in pages:
            props_str = json.dumps(props).lower() if isinstance(props, dict) else str(props).lower()
            if "earnings" in props_str or "insights" in props_str:
                found_earnings = True
                break
        check("Notion page mentions earnings or insights", found_earnings)

    check("At least 1 Notion block created", len(blocks) >= 1,
          f"Found {len(blocks)}")

    if blocks:
        block_text = " ".join(str(b[2] or "") for b in blocks).lower()
        has_revenue = "revenue" in block_text or "googl" in block_text or "amzn" in block_text
        check("Notion blocks mention revenue or companies", has_revenue)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("YF EARNINGS REPORT WORD NOTION - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace, gt_dir)
    check_notion()

    total_pass = PASS_COUNT
    total_fail = FAIL_COUNT
    all_ok = FAIL_COUNT == 0

    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": total_pass, "failed": total_fail, "success": all_ok}, f, indent=2)

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
