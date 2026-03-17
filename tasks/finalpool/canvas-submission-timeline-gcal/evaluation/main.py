"""
Evaluation script for canvas-submission-timeline-gcal task.

Checks:
1. Word file (Late_Submission_Report.docx) - contains 'late submission' and all Fall 2013 course codes
2. Google Calendar events >= 30 with course codes in summary
"""

import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

FALL_2013_CODES = ["AAA-2013J", "BBB-2013J", "DDD-2013J", "EEE-2013J", "FFF-2013J", "GGG-2013J"]

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail="", db=False):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        if db:
            PASS_COUNT += 1
        else:
            PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        if db:
            FAIL_COUNT += 1
        else:
            FAIL_COUNT += 1
        detail_str = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def check_word(agent_workspace):
    print("\n=== Checking Word Output ===")
    docx_path = os.path.join(agent_workspace, "Late_Submission_Report.docx")
    check("Word file exists", os.path.isfile(docx_path), f"Expected {docx_path}")
    if not os.path.isfile(docx_path):
        return False

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        check("Word file readable", False, str(e))
        return False

    # Collect all text
    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text.lower() + " "
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text.lower() + " "

    check("Word contains 'late submission'", "late submission" in all_text,
          f"Text sample: {all_text[:200]}")

    # Check all course codes present
    for code in FALL_2013_CODES:
        check(f"Word contains course code {code}",
              code.lower() in all_text,
              f"Not found in document")

    return True


def check_calendar():
    print("\n=== Checking Google Calendar ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT summary FROM gcal.events")
    events = cur.fetchall()
    cur.close()
    conn.close()

    print(f"[check_calendar] Found {len(events)} events.")

    # Count events that have a course code in summary
    course_events = []
    for (summary,) in events:
        if summary:
            s_lower = summary.lower()
            for code in FALL_2013_CODES:
                if code.lower() in s_lower:
                    course_events.append(summary)
                    break

    check("At least 30 calendar events with course codes",
          len(course_events) >= 30,
          f"Found {len(course_events)} events with course codes")

    # Check that at least 4 different course codes appear
    codes_found = set()
    for summary in course_events:
        for code in FALL_2013_CODES:
            if code.lower() in summary.lower():
                codes_found.add(code)
    check("Events cover at least 4 different course codes",
          len(codes_found) >= 4,
          f"Found codes: {codes_found}")

    return len(course_events) >= 30


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_calendar()

    total_pass = PASS_COUNT
    total_fail = FAIL_COUNT
    all_ok = FAIL_COUNT == 0

    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        result = {"passed": total_pass, "failed": total_fail, "success": all_ok}
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
